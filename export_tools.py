"""
============================================
🗺️ الخريطة: 07_export/export_tools.py
📌 الربط:
    - يستقبل من main.py (أوامر التصدير)
    - يستقبل من ai_core.py (المحتوى المعرفي)
    - يرسل روابط تحميل مباشرة
============================================
"""

# المتطلبات: reportlab, weasyprint, xlsxwriter, pdfkit

from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from weasyprint import HTML
import xlsxwriter
from docx import Document
from docx.shared import Inches, Pt
import pdfkit
import io
import base64
import json
from typing import Dict, Any, Union
import asyncio
from datetime import datetime

class DataExporter:
    """نظام التصدير الفوري للمستندات"""
    
    def __init__(self):
        self.status = "🟢 نشط"
        self.export_counter = 0
        print("🟢 Data Export - جاهز لتصدير أي تنسيق")
        
        # روابط التحميل المؤقتة
        self.download_links = {}
    
    async def export_to_pdf(self, content: Union[str, Dict], style: str = "professional") -> Dict:
        """تصدير المحادثة أو المحتوى إلى PDF"""
        try:
            if isinstance(content, dict):
                content = json.dumps(content, ensure_ascii=False, indent=2)
            
            # إنشاء PDF باستخدام ReportLab
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # إضافة عنوان
            title = Paragraph(f"تقرير Super-AI - {datetime.now().strftime('%Y-%m-%d %H:%M')}", 
                            styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # إضافة المحتوى
            for line in content.split('\n')[:100]:  # حد 100 سطر
                p = Paragraph(line, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 6))
            
            doc.build(story)
            
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            # إنشاء رابط تحميل
            download_id = f"pdf_{datetime.now().timestamp()}"
            pdf_base64 = base64.b64encode(pdf_bytes).decode()
            
            self.download_links[download_id] = {
                "data": pdf_base64,
                "format": "pdf",
                "created": datetime.now().isoformat()
            }
            
            return {
                "status": "success",
                "format": "PDF",
                "download_link": f"/api/v1/download/{download_id}",
                "file_size": len(pdf_bytes),
                "pages": len(content.split('\n')) // 30 + 1,
                "preview": pdf_base64[:50] + "..."
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"PDF خطأ في تصدير: {str(e)}"
            }
    
    async def export_to_word(self, content: Union[str, Dict]) -> Dict:
        """تصدير إلى Word"""
        try:
            if isinstance(content, dict):
                content = json.dumps(content, ensure_ascii=False, indent=2)
            
            doc = Document()
            
            # إضافة عنوان
            title = doc.add_heading(f'تقرير Super-AI', 0)
            title.alignment = 1  # توسيط
            
            # إضافة التاريخ
            doc.add_paragraph(f'تاريخ التصدير: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            
            # إضافة المحتوى
            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.style.font.size = Pt(11)
            
            # حفظ في الذاكرة
            buffer = io.BytesIO()
            doc.save(buffer)
            word_bytes = buffer.getvalue()
            buffer.close()
            
            download_id = f"docx_{datetime.now().timestamp()}"
            word_base64 = base64.b64encode(word_bytes).decode()
            
            self.download_links[download_id] = {
                "data": word_base64,
                "format": "docx",
                "created": datetime.now().isoformat()
            }
            
            return {
                "status": "success",
                "format": "Word",
                "download_link": f"/api/v1/download/{download_id}",
                "file_size": len(word_bytes),
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Word خطأ في تصدير: {str(e)}"
            }
    
    async def export_to_excel(self, data: Dict) -> Dict:
        """تصدير البيانات إلى Excel"""
        try:
            buffer = io.BytesIO()
            
            with xlsxwriter.Workbook(buffer) as workbook:
                # ورقة البيانات الرئيسية
                worksheet = workbook.add_worksheet("SuperAI Data")
                
                # تنسيقات
                header_format = workbook.add_format({
                    'bold': True,
                    'fg_color': '#4CAF50',
                    'font_color': 'white',
                    'border': 1
                })
                
                # كتابة الرؤوس
                headers = list(data.keys()) if isinstance(data, dict) else ["Content"]
                for col, header in enumerate(headers[:10]):  # حد 10 أعمدة
                    worksheet.write(0, col, header, header_format)
                
                # كتابة البيانات
                if isinstance(data, dict):
                    row = 1
                    for key, value in list(data.items())[:100]:  # حد 100 صف
                        worksheet.write(row, 0, str(key))
                        worksheet.write(row, 1, str(value)[:100])  # اختصار القيم الطويلة
                        row += 1
            
            excel_bytes = buffer.getvalue()
            buffer.close()
            
            download_id = f"xlsx_{datetime.now().timestamp()}"
            excel_base64 = base64.b64encode(excel_bytes).decode()
            
            self.download_links[download_id] = {
                "data": excel_base64,
                "format": "xlsx",
                "created": datetime.now().isoformat()
            }
            
            return {
                "status": "success",
                "format": "Excel",
                "download_link": f"/api/v1/download/{download_id}",
                "file_size": len(excel_bytes),
                "rows": min(len(data) if isinstance(data, dict) else 1, 100)
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Excel خطأ في تصدير: {str(e)}"
            }
    
    async def get_download_link(self, download_id: str) -> Dict:
        """استرجاع رابط التحميل"""
        if download_id in self.download_links:
            link_data = self.download_links[download_id]
            return {
                "status": "success",
                "data": link_data["data"],
                "format": link_data["format"],
                "filename": f"superai_export_{download_id}.{link_data['format']}"
            }
        return {
            "status": "error",
            "message": "رابط التحميل غير صالح أو منتهي الصلاحية"
        }
    
    async def export_conversation(self, conversation_id: str, cognitive_core) -> Dict:
        """تصدير محادثة كاملة من الذاكرة المعرفية"""
        if conversation_id in cognitive_core.conversation_memory:
            conv_data = cognitive_core.conversation_memory[conversation_id]
            return await self.export_to_pdf(conv_data)
        return {
            "status": "error",
            "message": "المحادثة غير موجودة"
        }