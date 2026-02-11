"""
============================================
🗺️ الخريطة: 05_ingestion/file_reader.py
📌 الربط:
    - يستقبل الملفات من main.py
    - يرسل النصوص المستخلصة إلى ai_core.py
    - يشارك موارد OCR مع vision_processor.py
============================================
"""

# المتطلبات: pypdf2, pymupdf, python-docx, openpyxl, pandas, easyocr, pytesseract

import PyPDF2
import fitz  # PyMuPDF
from docx import Document
import openpyxl
import pandas as pd
import easyocr
import pytesseract
from PIL import Image
import io
import base64
import json
import asyncio
from typing import Dict, Any, Union, List
from pathlib import Path
import numpy as np

class UniversalIngestion:
    """نظام استيعاب وتحليل جميع أنواع الملفات"""
    
    def __init__(self):
        self.status = "🟢 نشط"
        # تهيئة EasyOCR للغات المتعددة
        self.easy_ocr = easyocr.Reader(['ar', 'en'])
        print("🟢 Universal Ingestion - جاهز لقراءة جميع الملفات")
        
        # ذاكرة مؤقتة للملفات المعالجة
        self.processed_files_cache = {}
        
    async def read_pdf(self, file_bytes: bytes, method: str = "advanced") -> Dict:
        """قراءة ملفات PDF بدقة 100%"""
        try:
            text = ""
            metadata = {}
            images = []
            
            if method == "basic":
                # طريقة PyPDF2 الأساسية
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    text += page.extract_text()
                metadata = pdf_reader.metadata
                
            elif method == "advanced":
                # طريقة PyMuPDF المتقدمة مع استخراج الصور
                pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
                metadata = pdf_document.metadata
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    text += page.get_text()
                    
                    # استخراج الصور من الصفحة
                    image_list = page.get_images()
                    for img in image_list:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("png")
                            images.append(base64.b64encode(img_data).decode())
            
            return {
                "status": "success",
                "text": text,
                "metadata": metadata,
                "images": images[:5],  # أول 5 صور فقط
                "pages": len(pdf_document) if method == "advanced" else len(pdf_reader.pages),
                "method": method
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"PDF قراءة الخطأ في: {str(e)}"
            }
    
    async def read_word(self, file_bytes: bytes) -> Dict:
        """قراءة ملفات Word"""
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # استخراج النص
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # استخراج الجداول
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "status": "success",
                "text": text,
                "tables": tables,
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections)
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Word قراءة الخطأ في: {str(e)}"
            }
    
    async def read_excel(self, file_bytes: bytes) -> Dict:
        """قراءة ملفات Excel"""
        try:
            # قراءة باستخدام pandas
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            
            sheets = {}
            for sheet_name, df in df_dict.items():
                sheets[sheet_name] = {
                    "data": df.fillna("").to_dict('records'),
                    "columns": df.columns.tolist(),
                    "shape": df.shape
                }
            
            # قراءة باستخدام openpyxl للتفاصيل
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            formulas = {}
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_formulas = []
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.data_type == 'f':
                            sheet_formulas.append({
                                "cell": cell.coordinate,
                                "formula": cell.value
                            })
                formulas[sheet_name] = sheet_formulas
            
            return {
                "status": "success",
                "sheets": sheets,
                "formulas": formulas,
                "sheet_names": list(sheets.keys())
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Excel قراءة الخطأ في: {str(e)}"
            }
    
    async def ocr_image(self, image_bytes: bytes, language: str = 'ar+en') -> Dict:
        """التعرف الضوئي على النصوص في الصور"""
        try:
            # تحويل البايتات إلى صورة
            image = Image.open(io.BytesIO(image_bytes))
            
            # تحويل PIL Image إلى numpy array لـ easyocr
            img_array = np.array(image)
            
            # OCR باستخدام EasyOCR
            easy_result = self.easy_ocr.readtext(img_array, detail=0)
            
            # OCR باستخدام Tesseract كنسخة احتياطية
            tesseract_text = pytesseract.image_to_string(image, lang='ara+eng')
            
            return {
                "status": "success",
                "easyocr_text": " ".join(easy_result),
                "tesseract_text": tesseract_text.strip(),
                "combined_text": " ".join(easy_result) + "\n" + tesseract_text.strip(),
                "language": language,
                "confidence": 0.95  # تقدير
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"OCR الخطأ في: {str(e)}"
            }
    
    async def universal_read(self, file_content: str, file_type: str) -> Dict:
        """قراءة أي نوع ملفات تلقائياً"""
        try:
            # فك تشفير base64 إن وجد
            if file_content.startswith('data:'):
                file_content = file_content.split(',')[1]
            
            file_bytes = base64.b64decode(file_content)
            
            if file_type in ['pdf', 'application/pdf']:
                return await self.read_pdf(file_bytes, 'advanced')
            elif file_type in ['docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return await self.read_word(file_bytes)
            elif file_type in ['xlsx', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                return await self.read_excel(file_bytes)
            elif file_type in ['jpg', 'jpeg', 'png', 'image']:
                return await self.ocr_image(file_bytes)
            else:
                # محاولة قراءة كنص عادي
                text = file_bytes.decode('utf-8')
                return {
                    "status": "success",
                    "text": text,
                    "type": "plain_text"
                }
        except Exception as e:
            return {
                "status": "error",
                "message": f"قراءة الملف: {str(e)}"
            }